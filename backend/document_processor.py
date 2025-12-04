"""
Document Processor Module
Handles extraction of text from various document formats with line number tracking.
"""
import hashlib
import re
from pathlib import Path
from typing import Dict, List, Optional, Tuple
from dataclasses import dataclass, field
from datetime import datetime
import json

import pdfplumber
from docx import Document as DocxDocument
from tqdm import tqdm

from . import config
from .logger import log_info, log_success, log_step, log_document, log_error


@dataclass
class LineInfo:
    """Represents a single line with metadata."""
    line_number: int
    content: str
    page_number: Optional[int] = None
    paragraph_index: Optional[int] = None
    
    def to_dict(self) -> dict:
        return {
            "line_number": self.line_number,
            "content": self.content,
            "page_number": self.page_number,
            "paragraph_index": self.paragraph_index
        }


@dataclass
class DocumentChunk:
    """Represents a chunk of text with metadata for retrieval."""
    chunk_id: str
    content: str
    start_line: int
    end_line: int
    page_numbers: List[int]
    document_id: str
    document_name: str
    char_start: int
    char_end: int
    metadata: Dict = field(default_factory=dict)
    
    def to_dict(self) -> dict:
        return {
            "chunk_id": self.chunk_id,
            "content": self.content,
            "start_line": self.start_line,
            "end_line": self.end_line,
            "page_numbers": self.page_numbers,
            "document_id": self.document_id,
            "document_name": self.document_name,
            "char_start": self.char_start,
            "char_end": self.char_end,
            "metadata": self.metadata
        }


@dataclass
class ProcessedDocument:
    """Represents a fully processed document."""
    document_id: str
    name: str
    file_path: str
    file_type: str
    total_lines: int
    total_pages: int
    total_words: int
    total_characters: int
    lines: List[LineInfo]
    chunks: List[DocumentChunk]
    processed_at: str
    
    def to_dict(self) -> dict:
        return {
            "document_id": self.document_id,
            "name": self.name,
            "file_path": self.file_path,
            "file_type": self.file_type,
            "total_lines": self.total_lines,
            "total_pages": self.total_pages,
            "total_words": self.total_words,
            "total_characters": self.total_characters,
            "processed_at": self.processed_at
        }


class DocumentProcessor:
    """
    Processes various document formats and extracts text with line-level tracking.
    Supports PDF, Word (docx), and plain text files.
    """
    
    def __init__(
        self,
        chunk_size: int = config.CHUNK_SIZE,
        chunk_overlap: int = config.CHUNK_OVERLAP
    ):
        self.chunk_size = chunk_size
        self.chunk_overlap = chunk_overlap
        self._processors = {
            ".pdf": self._process_pdf,
            ".docx": self._process_docx,
            ".doc": self._process_docx,
            ".txt": self._process_text,
            ".md": self._process_text,
            ".csv": self._process_text,
        }
    
    def process_file(self, file_path: str) -> ProcessedDocument:
        """
        Process a document file and extract text with line numbers.
        
        Args:
            file_path: Path to the document file
            
        Returns:
            ProcessedDocument with all extracted content and metadata
        """
        path = Path(file_path)
        
        if not path.exists():
            raise FileNotFoundError(f"File not found: {file_path}")
        
        extension = path.suffix.lower()
        if extension not in self._processors:
            raise ValueError(
                f"Unsupported file type: {extension}. "
                f"Supported types: {list(self._processors.keys())}"
            )
        
        # Generate document ID
        doc_id = self._generate_document_id(path)
        
        # Process based on file type
        log_info(f"Processing {path.name}...")
        lines, page_count = self._processors[extension](path)
        
        # Calculate statistics
        total_words = sum(len(line.content.split()) for line in lines)
        total_chars = sum(len(line.content) for line in lines)
        
        log_step(f"Extracted {len(lines)} lines, {total_words:,} words")
        log_document("EXTRACTED", path.name, f"lines={len(lines)}, words={total_words}, pages={page_count}")
        
        # Create chunks with line tracking
        log_step("Creating chunks...")
        chunks = self._create_chunks(lines, doc_id, path.name)
        log_step(f"Created {len(chunks)} chunks")
        
        # Save processed document metadata
        processed_doc = ProcessedDocument(
            document_id=doc_id,
            name=path.name,
            file_path=str(path.absolute()),
            file_type=extension,
            total_lines=len(lines),
            total_pages=page_count,
            total_words=total_words,
            total_characters=total_chars,
            lines=lines,
            chunks=chunks,
            processed_at=datetime.now().isoformat()
        )
        
        # Save line mapping for future reference
        log_step("Saving document metadata...")
        self._save_line_mapping(processed_doc)
        log_step("Document processing complete")
        
        return processed_doc
    
    def _generate_document_id(self, path: Path) -> str:
        """Generate a unique document ID based on file content hash."""
        hasher = hashlib.md5()
        with open(path, 'rb') as f:
            for chunk in iter(lambda: f.read(8192), b''):
                hasher.update(chunk)
        return hasher.hexdigest()[:16]
    
    def _process_pdf(self, path: Path) -> Tuple[List[LineInfo], int]:
        """Extract text from PDF with page and line tracking."""
        lines = []
        line_number = 1
        page_count = 0
        
        with pdfplumber.open(path) as pdf:
            page_count = len(pdf.pages)
            log_step(f"Processing {page_count} pages...")
            
            for page_num, page in enumerate(pdf.pages, 1):
                text = page.extract_text() or ""
                
                for line_content in text.split('\n'):
                    if line_content.strip():  # Skip empty lines
                        lines.append(LineInfo(
                            line_number=line_number,
                            content=line_content.strip(),
                            page_number=page_num
                        ))
                        line_number += 1
        
        return lines, page_count
    
    def _process_docx(self, path: Path) -> Tuple[List[LineInfo], int]:
        """Extract text from Word document with paragraph tracking."""
        lines = []
        line_number = 1
        
        doc = DocxDocument(path)
        
        for para_idx, paragraph in enumerate(doc.paragraphs):
            text = paragraph.text.strip()
            if text:
                # Split long paragraphs into logical lines
                for line_content in text.split('\n'):
                    if line_content.strip():
                        lines.append(LineInfo(
                            line_number=line_number,
                            content=line_content.strip(),
                            paragraph_index=para_idx
                        ))
                        line_number += 1
        
        # Also extract from tables
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
                if row_text:
                    lines.append(LineInfo(
                        line_number=line_number,
                        content=row_text
                    ))
                    line_number += 1
        
        return lines, 1  # DOCX doesn't have pages in the same way
    
    def _process_text(self, path: Path) -> Tuple[List[LineInfo], int]:
        """Extract text from plain text files with line tracking."""
        lines = []
        
        # Try different encodings
        encodings = ['utf-8', 'utf-8-sig', 'latin-1', 'cp1252']
        content = None
        
        for encoding in encodings:
            try:
                with open(path, 'r', encoding=encoding) as f:
                    content = f.read()
                break
            except UnicodeDecodeError:
                continue
        
        if content is None:
            raise ValueError(f"Could not decode file: {path}")
        
        for line_number, line_content in enumerate(content.split('\n'), 1):
            if line_content.strip():
                lines.append(LineInfo(
                    line_number=line_number,
                    content=line_content.strip()
                ))
        
        return lines, 1
    
    def _create_chunks(
        self,
        lines: List[LineInfo],
        doc_id: str,
        doc_name: str
    ) -> List[DocumentChunk]:
        """
        Create overlapping chunks from lines while preserving line number mapping.
        """
        chunks = []
        
        if not lines:
            return chunks
        
        # Combine all lines into full text with line tracking
        full_text = ""
        line_char_positions = {}  # Maps character position to line number
        
        current_pos = 0
        for line in lines:
            line_char_positions[current_pos] = line.line_number
            full_text += line.content + "\n"
            current_pos = len(full_text)
        
        # Create chunks with overlap
        chunk_start = 0
        chunk_index = 0
        max_iterations = 10000  # Safety limit
        
        while chunk_start < len(full_text) and chunk_index < max_iterations:
            chunk_end = min(chunk_start + self.chunk_size, len(full_text))
            
            # Try to break at sentence or paragraph boundary
            if chunk_end < len(full_text):
                break_chars = ['\n\n', '.\n', '. ', '\n']
                for break_char in break_chars:
                    last_break = full_text.rfind(break_char, chunk_start, chunk_end)
                    if last_break > chunk_start + self.chunk_size // 2:
                        chunk_end = last_break + len(break_char)
                        break
            
            chunk_content = full_text[chunk_start:chunk_end].strip()
            
            if chunk_content:
                start_line = self._find_line_at_position(line_char_positions, chunk_start)
                end_line = self._find_line_at_position(line_char_positions, chunk_end)
                
                page_numbers = set()
                for line in lines:
                    if start_line <= line.line_number <= end_line:
                        if line.page_number:
                            page_numbers.add(line.page_number)
                
                chunk = DocumentChunk(
                    chunk_id=f"{doc_id}_{chunk_index}",
                    content=chunk_content,
                    start_line=start_line,
                    end_line=end_line,
                    page_numbers=sorted(page_numbers) if page_numbers else [],
                    document_id=doc_id,
                    document_name=doc_name,
                    char_start=chunk_start,
                    char_end=chunk_end,
                    metadata={"chunk_index": chunk_index, "total_chunks": -1}
                )
                chunks.append(chunk)
                chunk_index += 1
            
            # Move to next chunk - ensure progress
            new_start = chunk_end - self.chunk_overlap
            if new_start <= chunk_start:
                new_start = chunk_start + 100  # Force progress
            chunk_start = new_start
        
        # Update total chunks in metadata
        for chunk in chunks:
            chunk.metadata["total_chunks"] = len(chunks)
        
        return chunks
    
    def _find_line_at_position(
        self,
        line_char_positions: Dict[int, int],
        position: int
    ) -> int:
        """Find the line number at a given character position."""
        current_line = 1
        for char_pos, line_num in sorted(line_char_positions.items()):
            if char_pos <= position:
                current_line = line_num
            else:
                break
        return current_line
    
    def _save_line_mapping(self, doc: ProcessedDocument) -> None:
        """Save line mapping for quick lookup later."""
        mapping_file = config.PROCESSED_DIR / f"{doc.document_id}_lines.json"
        
        mapping = {
            "document_id": doc.document_id,
            "name": doc.name,
            "lines": {line.line_number: line.to_dict() for line in doc.lines}
        }
        
        with open(mapping_file, 'w', encoding='utf-8') as f:
            json.dump(mapping, f, ensure_ascii=False, indent=2)
    
    def get_lines_by_range(
        self,
        doc_id: str,
        start_line: int,
        end_line: int
    ) -> List[LineInfo]:
        """Retrieve specific lines from a processed document."""
        mapping_file = config.PROCESSED_DIR / f"{doc_id}_lines.json"
        
        if not mapping_file.exists():
            raise FileNotFoundError(f"Line mapping not found for document: {doc_id}")
        
        with open(mapping_file, 'r', encoding='utf-8') as f:
            mapping = json.load(f)
        
        lines = []
        for line_num in range(start_line, end_line + 1):
            line_data = mapping["lines"].get(str(line_num))
            if line_data:
                lines.append(LineInfo(**line_data))
        
        return lines
    
    def get_context_around_line(
        self,
        doc_id: str,
        line_number: int,
        context_lines: int = 3
    ) -> Tuple[List[LineInfo], int, int]:
        """
        Get lines around a specific line number for context.
        
        Returns:
            Tuple of (lines, start_line, end_line)
        """
        mapping_file = config.PROCESSED_DIR / f"{doc_id}_lines.json"
        
        with open(mapping_file, 'r', encoding='utf-8') as f:
            mapping = json.load(f)
        
        all_line_nums = sorted(int(k) for k in mapping["lines"].keys())
        max_line = max(all_line_nums) if all_line_nums else 0
        
        start = max(1, line_number - context_lines)
        end = min(max_line, line_number + context_lines)
        
        return self.get_lines_by_range(doc_id, start, end), start, end


class BatchProcessor:
    """Process multiple documents in batch."""
    
    def __init__(self, processor: Optional[DocumentProcessor] = None):
        self.processor = processor or DocumentProcessor()
    
    def process_directory(
        self,
        directory: str,
        recursive: bool = True
    ) -> List[ProcessedDocument]:
        """Process all supported documents in a directory."""
        path = Path(directory)
        documents = []
        
        pattern = "**/*" if recursive else "*"
        
        for file_path in path.glob(pattern):
            if file_path.is_file() and file_path.suffix.lower() in config.SUPPORTED_EXTENSIONS:
                try:
                    doc = self.processor.process_file(str(file_path))
                    documents.append(doc)
                except Exception as e:
                    print(f"⚠️ Error processing {file_path.name}: {e}")
        
        return documents
    
    def process_files(self, file_paths: List[str]) -> List[ProcessedDocument]:
        """Process a list of specific files."""
        documents = []
        
        for file_path in file_paths:
            try:
                doc = self.processor.process_file(file_path)
                documents.append(doc)
            except Exception as e:
                print(f"⚠️ Error processing {file_path}: {e}")
        
        return documents


if __name__ == "__main__":
    # Test the processor
    import sys
    
    if len(sys.argv) > 1:
        processor = DocumentProcessor()
        doc = processor.process_file(sys.argv[1])
        
        print(f"\n✅ Processed: {doc.name}")
        print(f"   Lines: {doc.total_lines:,}")
        print(f"   Words: {doc.total_words:,}")
        print(f"   Chunks: {len(doc.chunks)}")
        
        if doc.chunks:
            print(f"\n📝 Sample chunk (lines {doc.chunks[0].start_line}-{doc.chunks[0].end_line}):")
            print(doc.chunks[0].content[:500])

